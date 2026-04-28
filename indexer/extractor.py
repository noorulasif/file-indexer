"""
Text extraction module for the File Indexer.
Handles various file formats and extracts plain text content.
"""

import csv
import sys
import logging
from pathlib import Path
from typing import Optional

# Optional imports with fallbacks for missing dependencies
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

# Configure logging
logging.basicConfig(level=logging.WARNING, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

# Constants
MAX_TEXT_LENGTH = 3000
CSV_MAX_ROWS = 100


def extract_text(file_path: str) -> str:
    """
    Extract plain text content from a file.
    
    Args:
        file_path: Path to the file to extract text from
        
    Returns:
        Plain text content, truncated to MAX_TEXT_LENGTH characters.
        Returns empty string if extraction fails or file type is unsupported.
    """
    path = Path(file_path)
    
    # Check if file exists
    if not path.exists():
        logger.error(f"File not found: {file_path}")
        return ""
    
    # Check file size (optional, but good practice)
    try:
        if path.stat().st_size == 0:
            logger.warning(f"File is empty: {file_path}")
            return ""
    except OSError as e:
        logger.error(f"Cannot access file {file_path}: {e}")
        return ""
    
    # Get file extension and convert to lowercase
    extension = path.suffix.lower()
    
    # Route to appropriate extractor
    try:
        if extension in ['.txt', '.md']:
            text = _extract_text_file(path)
        elif extension == '.pdf':
            text = _extract_pdf(path)
        elif extension == '.docx':
            text = _extract_docx(path)
        elif extension == '.xlsx':
            text = _extract_xlsx(path)
        elif extension == '.csv':
            text = _extract_csv(path)
        elif extension in ['.jpg', '.jpeg', '.png']:
            text = "IMAGE_FILE"
        else:
            # Unsupported file type
            logger.debug(f"Unsupported file type: {extension}")
            return ""
        
        # Truncate if necessary
        if len(text) > MAX_TEXT_LENGTH:
            original_len = len(text)
            text = text[:MAX_TEXT_LENGTH]
            logger.debug(f"Truncated text from {original_len} to {MAX_TEXT_LENGTH} chars")
        
        return text
        
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        return ""


def _extract_text_file(path: Path) -> str:
    """
    Extract text from plain text or markdown files.
    
    Args:
        path: Path to the text file
        
    Returns:
        File contents as string
    """
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Error reading text file {path} with {encoding}: {e}")
            continue
    
    # If all encodings fail
    logger.error(f"Could not decode {path} with any supported encoding")
    return ""


def _extract_pdf(path: Path) -> str:
    """
    Extract text from PDF using pdfplumber.
    
    Args:
        path: Path to the PDF file
        
    Returns:
        Extracted text from all pages
    """
    if pdfplumber is None:
        logger.error("pdfplumber not installed. Install with: pip install pdfplumber")
        return ""
    
    try:
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
                else:
                    logger.debug(f"No text found on page {page_num} of {path}")
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"PDF extraction error for {path}: {e}")
        return ""


def _extract_docx(path: Path) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        path: Path to the DOCX file
        
    Returns:
        Concatenated text from all paragraphs
    """
    if Document is None:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return ""
    
    try:
        doc = Document(path)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n".join(paragraphs)
    
    except Exception as e:
        logger.error(f"DOCX extraction error for {path}: {e}")
        return ""


def _extract_xlsx(path: Path) -> str:
    """
    Extract text from XLSX file.
    
    Args:
        path: Path to the XLSX file
        
    Returns:
        Text representation of cell values in format "Sheet: SheetName\nRow X: values"
    """
    if load_workbook is None:
        logger.error("openpyxl not installed. Install with: pip install openpyxl")
        return ""
    
    try:
        workbook = load_workbook(path, data_only=True)
        all_sheets = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_content = [f"Sheet: {sheet_name}"]
            
            rows_processed = 0
            for row in sheet.iter_rows(values_only=True):
                # Skip completely empty rows
                if all(cell is None for cell in row):
                    continue
                
                # Convert row to text
                row_values = []
                for col_idx, cell in enumerate(row, 1):
                    if cell is not None:
                        # Convert to string, handling different types
                        cell_str = str(cell).strip()
                        if cell_str:
                            row_values.append(f"Col{col_idx}: {cell_str}")
                
                if row_values:
                    sheet_content.append(" | ".join(row_values))
                    rows_processed += 1
                
                # Limit rows to avoid huge output
                if rows_processed >= 200:
                    sheet_content.append(f"... (truncated after {rows_processed} rows)")
                    break
            
            all_sheets.append("\n".join(sheet_content))
        
        return "\n\n".join(all_sheets)
    
    except Exception as e:
        logger.error(f"XLSX extraction error for {path}: {e}")
        return ""


def _extract_csv(path: Path) -> str:
    """
    Extract text from CSV file.
    
    Args:
        path: Path to the CSV file
        
    Returns:
        Text representation of first CSV_MAX_ROWS rows
    """
    try:
        rows = []
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            
            for idx, row in enumerate(reader):
                if idx >= CSV_MAX_ROWS:
                    rows.append(f"... (truncated after {CSV_MAX_ROWS} rows)")
                    break
                
                # Clean and format the row
                clean_row = [cell.strip() if cell else "" for cell in row]
                rows.append(", ".join(clean_row))
        
        if not rows:
            return ""
        
        return "\n".join(rows)
    
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            rows = []
            with open(path, 'r', encoding='latin-1') as f:
                reader = csv.reader(f)
                for idx, row in enumerate(reader):
                    if idx >= CSV_MAX_ROWS:
                        rows.append(f"... (truncated after {CSV_MAX_ROWS} rows)")
                        break
                    clean_row = [cell.strip() if cell else "" for cell in row]
                    rows.append(", ".join(clean_row))
            return "\n".join(rows)
        except Exception as e:
            logger.error(f"CSV encoding error for {path}: {e}")
            return ""
    
    except Exception as e:
        logger.error(f"CSV extraction error for {path}: {e}")
        return ""


def main():
    """Test the extractor from command line."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Extract text from various file formats"
    )
    parser.add_argument(
        "file_path",
        help="Path to the file to extract text from"
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Show verbose output (including logging)"
    )
    parser.add_argument(
        "--no-truncate",
        action="store_true",
        help="Show full text without truncation"
    )
    
    args = parser.parse_args()
    
    # Configure logging level
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    
    # Check if file exists
    file_path = Path(args.file_path)
    if not file_path.exists():
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
    
    # Extract text
    print(f"Extracting text from: {file_path}")
    print(f"File size: {file_path.stat().st_size:,} bytes")
    print(f"File type: {file_path.suffix}")
    print("-" * 60)
    
    # Temporarily remove truncation for testing if requested
    global MAX_TEXT_LENGTH
    original_limit = MAX_TEXT_LENGTH
    if args.no_truncate:
        import builtins
        builtins.MAX_TEXT_LENGTH = float('inf')
    
    try:
        text = extract_text(str(file_path))
        
        if not text:
            print("No text extracted (empty or unsupported file)")
        elif text == "IMAGE_FILE":
            print("IMAGE_FILE - This is an image file (will be processed by vision model)")
        else:
            print(f"Extracted text ({len(text)} characters):")
            print("-" * 60)
            print(text)
            print("-" * 60)
            
            if len(text) >= MAX_TEXT_LENGTH and not args.no_truncate:
                print(f"\n⚠️  Text truncated to {MAX_TEXT_LENGTH} characters")
                print("   Use --no-truncate to show full content")
    
    except Exception as e:
        print(f"Error during extraction: {e}")
        sys.exit(1)
    
    finally:
        # Restore original limit
        if args.no_truncate:
            builtins.MAX_TEXT_LENGTH = original_limit


if __name__ == "__main__":
    main()



